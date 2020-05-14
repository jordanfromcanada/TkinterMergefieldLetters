from fuzzywuzzy import fuzz, process
from mailmerge import MailMerge
from docx.shared import Inches
from datetime import date
from docx import Document
import comtypes.client
from tkinter import *
import functools
import pprint
import os

# Word template filepath
letter_template = r'template.docx'

# Return available MailMerge fields within the given Word file
def checkfields(letter_template_path)
    document = MailMerge(letter_template_path)
    return [i for i in document.get_merge_fields()]
    
# Inputs mailmerge fields retrived from the user and saves the output as both Word and PDF files
def makeLetter(letter_template_path, filled_fields_dict)
    todays_date = date.today().strftime(%b-%d-%Y) #e.g. Feb-23-2020
    # note: document object is of class 'mailmerge.MailMerge'
    document = MailMerge(letter_template_path)
    document.merge_pages([filled_fields_dict])
    output_dir = r'output-folder'
    document_outpath = os.path.join(output_dir, 'word_file_' + todays_date + '.docx')
    document.write(document_outpath)
    document.close()

    # Ensure a table exists in the word file in order to insert the .png signature image within the table
    doc = Document(document_outpath)
    tables = doc.tables
    assert tables, 'You need to insert a 1x1 empty table under the Signature line in the letter template!'
    p = tables[0].rows[0].cells[0].add_paragraph()
    r = p.add_run()

    # Find the correct signature image filename using fuzzy string matching
    signatures_folder = r'signatures-folder'
    # For example purposes, the letter is signed by the manager 'William H Taft' (which retrieves the signature William H Taft.png)
    highest_name_match = process.extractOne('William H Taft', os.listdir(signatures_folder))
    print('Best signature filename match to your entry=')
    print(os.listdir(signatures_folder), highest_name_match)

    # Insert the signature image into the table and save the word doc with the same name
    r.add_picture(os.path.join(signatures_folder, highest_name_match[0]), height=Inches(.4))
    doc.save(document_outpath)

    # Convert the word file to PDF and save it in the same directory
    wdFormatPDF = 17
    in_file = os.path.join(document_outpath)
    out_file = os.path.join(output_dir, 'word_to_PDF_' + todays_date + '.pdf')
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

# Create the class containing the Tkinter GUI
class App
    def __init__(self, root)
        self.root = root
        self.frame = Frame(self.root)
        self.frame.pack()
        
        # Create the initial form selection combobox and specify available letters
        self.combo_box_selection = Combobox(self.root,
                                    values=[
                                            'Letter 1',
                                            'Letter 2',
                                            'Letter 3',
                                            'Letter 4'])
        self.combo_box_selection.pack(side=TOP)
        self.combo_box_selection.bind('<<ComboboxSelected>>', self.command)
        
    # Used to clear the entry field text when the field is clicked
    def clear_search(self, event, e)
        e.delete(0, END)
        
    # When ComboBox letter is selected, a new frame is created with all available mergefields as Entry boxes
    def command(self, event=None)
        for widget in self.frame.winfo_children()
            print('widget=', widget) # Debugging; duplicate frames causing duplicate entrybox texts
            widget.destroy()
        self.frame.destroy() # Destroys frame and contents
        self.frame = Frame(self.root) # Recreates frame
        self.frame.pack(pady=5, padx=5)
        self.text = [] # Empty array to store entry boxes
        self.labels = [] # Empty array to store entry boxes
        
        # To make extensible to multiple .docx templates in the same directory, use
        # extensible_templates = {'Letter 1' 'filepath to letter 1 .docx', 'Letter 2' 'filepath to letter 1 .docx'...}
        
        # Obtain available MailMerge fieldnames as a list
        self.word_doc_fields = checkfields(letter_template)
        self.letter_selection = self.combo_box_selection.get()
        print(self.letter_selection, self.word_doc_fields) # Print the letter type selected (e.g. Letter 1)
        
        # Create labelled input boxes for each MailMerge fieldname from the Word doc
        for i, j in enumerate(self.word_doc_fields)
            # Use this to display descriptive labels over each entry box
            # self.labels.append(Label(self.frame, width=15, text=Label Question +str(i), anchor='w'))
            # self.labels[i].pack()

            self.text.append(Entry(self.frame, text=str(j)))
            self.text[i].insert(0, str(j))
            print(self.text[i].get())
            self.text[i].pack()
            self.text[i].bind('<<Button-1>>', functools.partial(self.clear_search, e=self.text[i]))
        #print('i=', i, 'j=', j)
        #print(self.text)
        
        # Upon buttom click, fieldnames and entries are stored in a dictionary 'filled_fields_dict'
        self.done = Button(self.frame, text='Create my letter in Word and PDF!', command=self.fields_as_dict)
        self.done.pack()
        self.frame.pack()
        
    # Retrieves inputted user entries and calls makeLetter
    def fields_as_dict(self)
        self.dict = {}
        for i in range(len(self.text))
            # Adds a new dict entry in the format {'Entry fieldname' 'User input'}
            self.dict.update({self.text[i].cget('text') self.text[i].get()})
        # Prints dict of entries
        print('Entry data from {}:\n'.format(self.letter_selection))
        pp = pprint.PrettyPrinter(indent=4)
        print('Fields=\n')
        pp.pprint(self.dict) #prints the dict
        global filled_fields_dict # made global for testing purposes in jupyter, can remove later
        filled_fields_dict = self.dict
        # This creates the filled Word and PDF files in the 'output_dir' directory
        makeLetter(letter_template, filled_fields_dict)
        
root = Tk()
App(root)
root.mainloop()
